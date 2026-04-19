
# main.py
import time
import os
from io import BytesIO
from datetime import datetime

import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from descriptors import DESCRIPTORS


# =========================
# CUSTOM CSS + UI HELPERS
# =========================

def inject_css():
    st.markdown("""
    <style>
    .step-track{display:flex;gap:0;margin-bottom:20px;border:1px solid #e5e7eb;border-radius:10px;overflow:hidden}
    .step{flex:1;padding:10px 8px;text-align:center;border-right:1px solid #e5e7eb;background:#f9fafb}
    .step:last-child{border-right:none}
    .step.s-done{background:#f0fdf4}.step.s-active{background:#ede9fe}.step.s-locked{opacity:.6}
    .step-num{display:inline-flex;width:20px;height:20px;border-radius:50%;align-items:center;justify-content:center;font-size:10px;font-weight:700;margin-bottom:3px}
    .step.s-done .step-num{background:#bbf7d0;color:#166534}
    .step.s-active .step-num{background:#c4b5fd;color:#3730a3}
    .step.s-locked .step-num{background:#e5e7eb;color:#9ca3af}
    .step-lbl{font-size:11px;line-height:1.3;color:#6b7280}
    .step.s-done .step-lbl{color:#166534}.step.s-active .step-lbl{color:#4c1d95;font-weight:500}
    .guidance-box{background:#ede9fe;border:1px solid #c4b5fd;border-radius:8px;padding:12px 14px;margin-bottom:16px;font-size:13px;color:#3730a3;line-height:1.6}
    .guidance-box .g-lbl{font-size:10px;font-weight:700;letter-spacing:.8px;text-transform:uppercase;color:#6d28d9;margin-bottom:4px}
    .metrics-row{display:flex;gap:10px;margin-bottom:18px;flex-wrap:wrap}
    .metric-card{flex:1;min-width:110px;padding:12px 14px;border:1px solid #e5e7eb;border-radius:10px;background:#fff;text-align:center}
    .metric-card .mc-val{font-size:26px;font-weight:700;line-height:1;margin-bottom:3px}
    .metric-card .mc-lbl{font-size:11px;color:#6b7280;line-height:1.3}
    .mc-green{background:#f0fdf4;border-color:#bbf7d0}.mc-green .mc-val{color:#166534}
    .mc-amber{background:#fffbeb;border-color:#fde68a}.mc-amber .mc-val{color:#92400e}
    .rating-pill{display:inline-block;padding:2px 7px;border-radius:4px;font-size:11px;font-weight:700}
    .rp-he{background:#bbf7d0;color:#166534}.rp-e{background:#bfdbfe;color:#1e3a8a}
    .rp-in{background:#fef08a;color:#713f12}.rp-dn{background:#fecaca;color:#7f1d1d}
    .cmp-wrap{margin-bottom:14px}
    .cmp-domain-lbl{font-size:11px;font-weight:600;color:#6b7280;text-transform:uppercase;letter-spacing:.6px;padding:5px 9px;background:#f9fafb;border:1px solid #e5e7eb;border-radius:6px 6px 0 0;border-bottom:none}
    .cmp-table{width:100%;border-collapse:collapse}
    .cmp-table th{padding:6px 9px;background:#f3f4f6;font-size:11px;font-weight:600;color:#374151;border:1px solid #e5e7eb;text-align:left}
    .cmp-table td{padding:5px 9px;border:1px solid #e5e7eb;font-size:12px;vertical-align:middle}
    .trend-up{color:#166534;font-weight:600;font-size:12px}
    .trend-dn{color:#7f1d1d;font-weight:600;font-size:12px}
    .trend-nc{color:#9ca3af;font-size:12px}.trend-na{color:#d1d5db;font-size:12px}
    .sec-heading{font-size:14px;font-weight:600;color:#111827;margin:0 0 10px;padding-bottom:6px;border-bottom:1px solid #e5e7eb}
    </style>
    """, unsafe_allow_html=True)


def render_step_tracker(steps):
    parts = []
    for i, (label, subtitle, state) in enumerate(steps, 1):
        icon = "✓" if state == "done" else str(i)
        parts.append(
            f'<div class="step s-{state}">' +
            f'<div class="step-num">{icon}</div>' +
            f'<div class="step-lbl"><strong>{label}</strong><br>{subtitle}</div>' +
            '</div>'
        )
    st.markdown('<div class="step-track">' + "".join(parts) + '</div>', unsafe_allow_html=True)


def render_guidance(title, body):
    st.markdown(
        f'<div class="guidance-box"><div class="g-lbl">{title}</div>{body}</div>',
        unsafe_allow_html=True
    )


def render_metrics(metrics):
    cards = ""
    for val, label, style in metrics:
        cls = f"metric-card mc-{style}" if style else "metric-card"
        cards += f'<div class="{cls}"><div class="mc-val">{val}</div><div class="mc-lbl">{label}</div></div>'
    st.markdown(f'<div class="metrics-row">{cards}</div>', unsafe_allow_html=True)


def rating_pill_html(rating):
    short = rating_short(rating)
    cls_map = {"HE": "rp-he", "E": "rp-e", "IN": "rp-in", "DNMS": "rp-dn"}
    cls = cls_map.get(short, "")
    if not short:
        return '<span style="color:#d1d5db">—</span>'
    return f'<span class="rating-pill {cls}">{short}</span>'


def trend_html_cell(init_val, final_val):
    if not init_val or not final_val:
        return '<span class="trend-na">—</span>'
    arrow = trend_arrow(init_val, final_val)
    if "Improved" in arrow:
        return '<span class="trend-up">↑ improved</span>'
    elif "Dropped" in arrow:
        return '<span class="trend-dn">↓ dropped</span>'
    elif "No change" in arrow:
        return '<span class="trend-nc">→ same</span>'
    return '<span class="trend-na">—</span>'


def render_comparison_table_new(comparison_df):
    domain_titles = {
        "A": "Domain A — Planning and Preparation for Learning",
        "B": "Domain B — Classroom Management",
        "C": "Domain C — Delivery of Instruction",
        "D": "Domain D — Monitoring, Assessment, and Follow-Up",
        "E": "Domain E — Family and Community Outreach",
        "F": "Domain F — Professional Responsibility",
    }
    if comparison_df.empty:
        st.info("No comparison data available yet.")
        return
    for domain_key in ["A", "B", "C", "D", "E", "F"]:
        dom_df = comparison_df[comparison_df["Domain"] == domain_key]
        if dom_df.empty:
            continue
        rows_html = ""
        for _, row in dom_df.iterrows():
            init_v = safe_text(row.get("Initial", ""))
            final_v = safe_text(row.get("Final", ""))
            rows_html += (
                "<tr>"
                f'<td style="width:45%">{safe_text(row.get("Strand",""))}</td>' +
                f'<td style="width:18%;text-align:center">{rating_pill_html(init_v)}</td>' +
                f'<td style="width:18%;text-align:center">{rating_pill_html(final_v)}</td>' +
                f'<td style="width:19%">{trend_html_cell(init_v, final_v)}</td>' +
                "</tr>"
            )
        label = domain_titles.get(domain_key, f"Domain {domain_key}")
        st.markdown(
            f'<div class="cmp-wrap">' +
            f'<div class="cmp-domain-lbl">{label}</div>' +
            '<table class="cmp-table"><thead><tr>' +
            '<th>Strand</th><th style="text-align:center">Initial</th><th style="text-align:center">Final</th><th>Trend</th>' +
            f'</tr></thead><tbody>{rows_html}</tbody></table></div>',
            unsafe_allow_html=True
        )


# =========================
# Helper: add descriptors as subheaders (inline under column names)
# =========================

def add_descriptor_subheaders(df):
    """
    Append short Kim Marshall descriptors under each rubric column header.
    Uses HE (Highly Effective) summary line for quick context.
    """
    new_cols = []
    for col in df.columns:
        code = col.split()[0] if " " in col else col
        if code in DESCRIPTORS:
            short_desc = DESCRIPTORS[code]["HE"]
            if len(short_desc) > 80:  # truncate long ones
                short_desc = short_desc[:77] + "..."
            new_cols.append(f"{col}\n🛈 {short_desc}")
        else:
            new_cols.append(col)
    df.columns = new_cols
    return df

def safe_text(value):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return str(value)

def title_case_name(name: str) -> str:
    return " ".join(part.capitalize() for part in safe_text(name).split())

def highlight_ratings(val):
    colors = {
        "HE": "background-color: #a8e6a1;",   # green
        "E": "background-color: #d0f0fd;",    # blue
        "IN": "background-color: #fff3b0;",   # yellow
        "DNMS": "background-color: #f8a5a5;", # red
        "Highly Effective": "background-color: #a8e6a1;",
        "Effective": "background-color: #d0f0fd;",
        "Improvement Necessary": "background-color: #fff3b0;",
        "Does Not Meet Standards": "background-color: #f8a5a5;",
    }
    return colors.get(val, "")
 
def rating_rank(value):
    order = {
        "Does Not Meet Standards": 1,
        "Improvement Necessary": 2,
        "Effective": 3,
        "Highly Effective": 4,
        "DNMS": 1,
        "IN": 2,
        "E": 3,
        "HE": 4,
    }
    return order.get(str(value).strip(), 0)


def rating_short(value):
    mapping = {
        "Highly Effective": "HE",
        "Effective": "E",
        "Improvement Necessary": "IN",
        "Does Not Meet Standards": "DNMS",
        "HE": "HE",
        "E": "E",
        "IN": "IN",
        "DNMS": "DNMS",
    }
    return mapping.get(str(value).strip(), str(value).strip())


def trend_arrow(initial_value, final_value):
    init_score = rating_rank(initial_value)
    final_score = rating_rank(final_value)

    if init_score == 0 or final_score == 0:
        return ""
    if final_score > init_score:
        return "↑ Improved"
    if final_score < init_score:
        return "↓ Dropped"
    return "→ No change"


def trend_style(val):
    styles = {
        "↑ Improved": "background-color: #d9f2d9; color: #1f6f1f; font-weight: bold;",
        "↓ Dropped": "background-color: #f8d7da; color: #842029; font-weight: bold;",
        "→ No change": "background-color: #eef2f7; color: #495057;"
    }
    return styles.get(val, "")

def build_initial_final_comparison(rows_df):
    """
    Returns:
      - latest_initial
      - latest_final
      - comparison_df (vertical, good for appraisers)
    """
    if rows_df.empty:
        return None, None, pd.DataFrame()

    working = rows_df.copy()

    if "Assessment Cycle" not in working.columns:
        working["Assessment Cycle"] = "Initial"
    else:
        working["Assessment Cycle"] = working["Assessment Cycle"].replace("", "Initial")

    initial_rows = working[working["Assessment Cycle"] == "Initial"]
    final_rows = working[working["Assessment Cycle"] == "Final"]

    latest_initial = (
        initial_rows.sort_values("Timestamp", ascending=False).head(1)
        if not initial_rows.empty else None
    )
    latest_final = (
        final_rows.sort_values("Timestamp", ascending=False).head(1)
        if not final_rows.empty else None
    )

    comparison_rows = []

    for domain, items in DOMAINS.items():
        for code, label in items:
            strand = f"{code} {label}"

            init_val = ""
            final_val = ""

            if latest_initial is not None and not latest_initial.empty:
                init_val = safe_text(latest_initial.iloc[0].get(strand, ""))

            if latest_final is not None and not latest_final.empty:
                final_val = safe_text(latest_final.iloc[0].get(strand, ""))

            changed = "Yes" if init_val != final_val else "No"

            comparison_rows.append({
                "Domain": domain.split(":")[0],   # just A / B / C / D / E / F
                "Strand": strand,
                "Explanation": DESCRIPTORS.get(strand, {}).get("HE", ""),
                "Initial": rating_short(init_val),
                "Final": rating_short(final_val),
                "Trend": trend_arrow(init_val, final_val),
            })

    comparison_df = pd.DataFrame(comparison_rows)
    return latest_initial, latest_final, comparison_df

def rating_to_descriptor_key(rating_text):
    mapping = {
        "Highly Effective": "HE",
        "Effective": "E",
        "Improvement Necessary": "IN",
        "Does Not Meet Standards": "DNMS",
        "HE": "HE",
        "E": "E",
        "IN": "IN",
        "DNMS": "DNMS",
    }
    return mapping.get(safe_text(rating_text), "")
    
def add_summary_section_to_doc(doc, latest_record):
    """
    Creates a clean appraiser-friendly summary only.
    No rubric pages, no template content.
    Includes strand code/name, selected rating, and descriptor explanation.
    """
    doc.add_heading("OIS Teacher Appraisal Summary", level=1)

    p = doc.add_paragraph()
    p.add_run("Teacher: ").bold = True
    p.add_run(safe_text(latest_record.get("Name", "")))

    p = doc.add_paragraph()
    p.add_run("Appraiser: ").bold = True
    p.add_run(safe_text(latest_record.get("Appraiser", "")))

    p = doc.add_paragraph()
    p.add_run("Submitted on: ").bold = True
    p.add_run(safe_text(latest_record.get("Timestamp", "")))

    p = doc.add_paragraph()
    p.add_run("Last edited on: ").bold = True
    p.add_run(safe_text(latest_record.get("Last Edited On", "")))

    doc.add_paragraph("")

    for domain, items in DOMAINS.items():
        doc.add_heading(domain, level=2)

        domain_reflection = safe_text(latest_record.get(f"{domain} Reflection", ""))

        for code, label in items:
            strand_key = f"{code} {label}"
            selected_rating = safe_text(latest_record.get(strand_key, ""))
            descriptor_key = rating_to_descriptor_key(selected_rating)

            explanation = ""
            if strand_key in DESCRIPTORS and descriptor_key in DESCRIPTORS[strand_key]:
                explanation = safe_text(DESCRIPTORS[strand_key][descriptor_key])

            p = doc.add_paragraph()
            p.add_run(f"{strand_key}\n").bold = True
            p.add_run("Selected Rating: ").bold = True
            p.add_run(f"{selected_rating}\n")
            p.add_run("Explanation: ").bold = True
            p.add_run(explanation if explanation else "No explanation found.")

        if domain_reflection:
            p = doc.add_paragraph()
            p.add_run("Domain Reflection: ").bold = True
            p.add_run(domain_reflection)

        doc.add_paragraph("")


def generate_teacher_docx(teacher_name, latest_df):
    """
    Generates a clean DOCX with summary only.
    Does not use the rubric template.
    """
    latest_record = latest_df.iloc[0].to_dict()

    doc = Document()
    add_summary_section_to_doc(doc, latest_record)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def generate_final_evaluation_docx(record: dict):
    template_path = os.path.join(os.path.dirname(__file__), "..", "Copy of Letter template OIS JVLR.docx")

    try:
        doc = Document(template_path)
    except Exception:
        doc = Document()

    teacher_name = title_case_name(record.get("Teacher Name", ""))
    appraiser_name = title_case_name(record.get("Appraiser", ""))
    subject_area = safe_text(record.get("Subject Area", ""))

    # ===== TITLE =====
    # Remove extra spacing at top (important)
    if doc.paragraphs:
        first_para = doc.paragraphs[0]
        first_para.paragraph_format.space_before = Pt(0)
        first_para.paragraph_format.space_after = Pt(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    
    run = p.add_run("FINAL EVALUATION SUMMARY")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    # ===== BASIC DETAILS =====
    p = doc.add_paragraph()
    p.add_run("Teacher: ").bold = True
    p.add_run(teacher_name)

    p = doc.add_paragraph()
    p.add_run("Appraiser: ").bold = True
    p.add_run(appraiser_name)

    p = doc.add_paragraph()
    p.add_run("Subject Area: ").bold = True
    p.add_run(subject_area)

    doc.add_paragraph("")

    # ===== STUDENT SURVEY =====
    p = doc.add_paragraph()
    run = p.add_run("Student Survey Feedback")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("Administered by the teacher each Semester").italic = True

    doc.add_paragraph(safe_text(record.get("Student Survey Feedback", "")))
    doc.add_paragraph("")

    # ===== TEACHER REFLECTION =====
    p = doc.add_paragraph()
    run = p.add_run("Overall Reflection by the teacher on the school year")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(safe_text(record.get("Overall Reflection", "")))
    doc.add_paragraph("")

    # ===== DOMAIN RATINGS =====
    p = doc.add_paragraph()
    run = p.add_run("Ratings on Individual Rubrics")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("Domain").bold = True
    p.add_run(" | ").bold = True
    p.add_run("Rating").bold = True

    for col_name, label in final_eval_domain_rows():
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(safe_text(record.get(col_name, "")))

    doc.add_paragraph("")

    # ===== OVERALL RATING =====
    p = doc.add_paragraph()
    run = p.add_run("Overall Rating")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(safe_text(record.get("Overall Rating", "")))
    run.bold = True

    doc.add_paragraph("")

    # ===== APPRAISER COMMENTS =====
    p = doc.add_paragraph()
    run = p.add_run("Overall Appraiser Comments")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(safe_text(record.get("Overall Comments", "")))
    doc.add_paragraph("")

    # ===== SIGN OFF =====
    p = doc.add_paragraph()
    run = p.add_run("Sign Off")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run(f"{appraiser_name} signed off on: ").bold = True
    p.add_run(safe_text(record.get("Evaluator Sign Off Date", "")))

    p = doc.add_paragraph()
    p.add_run(f"{teacher_name} signed off on: ").bold = True
    p.add_run(safe_text(record.get("Teacher Sign Off Date", "")))

    doc.add_paragraph("")
    doc.add_paragraph(
        "The teacher’s signature indicates that he or she has seen and discussed the evaluation; "
        "it does not necessarily denote agreement with the report."
    )

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def build_teacher_initial_final(email):
    df = load_responses_df()
    mine = df[df["Email"] == email.strip().lower()] if not df.empty else pd.DataFrame()

    if mine.empty:
        return None, None, pd.DataFrame()

    if "Assessment Cycle" not in mine.columns:
        mine["Assessment Cycle"] = "Initial"
    else:
        mine["Assessment Cycle"] = mine["Assessment Cycle"].replace("", "Initial")

    initial_rows = mine[mine["Assessment Cycle"] == "Initial"]
    final_rows = mine[mine["Assessment Cycle"] == "Final"]

    latest_initial = (
        initial_rows.sort_values("Timestamp", ascending=False).head(1)
        if not initial_rows.empty else None
    )
    latest_final = (
        final_rows.sort_values("Timestamp", ascending=False).head(1)
        if not final_rows.empty else None
    )

    comparison_rows = []

    for domain, items in DOMAINS.items():
        for code, label in items:
            strand = f"{code} {label}"

            init_val = ""
            final_val = ""

            if latest_initial is not None and not latest_initial.empty:
                init_val = safe_text(latest_initial.iloc[0].get(strand, ""))

            if latest_final is not None and not latest_final.empty:
                final_val = safe_text(latest_final.iloc[0].get(strand, ""))

            comparison_rows.append({
                "Domain": domain.split(":")[0],
                "Strand": strand,
                "Explanation": DESCRIPTORS.get(strand, {}).get("HE", ""),
                "Initial": rating_short(init_val),
                "Final": rating_short(final_val),
                "Trend": trend_arrow(init_val, final_val),
            })

    comparison_df = pd.DataFrame(comparison_rows)
    return latest_initial, latest_final, comparison_df

def render_comparison_html(df):
    if df.empty:
        return "<p>No comparison data available.</p>"

    rating_bg = {
        "HE": "#a8e6a1",
        "E": "#d0f0fd",
        "IN": "#fff3b0",
        "DNMS": "#f8a5a5",
    }

    trend_bg = {
        "↑ Improved": "#d9f2d9",
        "↓ Dropped": "#f8d7da",
        "→ No change": "#eef2f7",
        "": "#ffffff",
    }

    html = """
    <div style="overflow-x:auto;">
      <table style="
          border-collapse: collapse;
          width: 100%;
          table-layout: fixed;
          font-family: Arial, sans-serif;
          font-size: 13px;
      ">
        <thead>
          <tr style="background-color:#f5f6f7;">
            <th style="border:1px solid #ddd; padding:8px; width:7%; text-align:left;">Domain</th>
            <th style="border:1px solid #ddd; padding:8px; width:14%; text-align:left;">Strand</th>
            <th style="border:1px solid #ddd; padding:8px; width:49%; text-align:left;">Explanation</th>
            <th style="border:1px solid #ddd; padding:8px; width:8%; text-align:center;">Initial</th>
            <th style="border:1px solid #ddd; padding:8px; width:8%; text-align:center;">Final</th>
            <th style="border:1px solid #ddd; padding:8px; width:14%; text-align:center;">Trend</th>
          </tr>
        </thead>
        <tbody>
    """

    for _, row in df.iterrows():
        initial = safe_text(row.get("Initial", ""))
        final = safe_text(row.get("Final", ""))
        trend = safe_text(row.get("Trend", ""))

        initial_bg = rating_bg.get(initial, "#ffffff")
        final_bg = rating_bg.get(final, "#ffffff")
        trend_bg_color = trend_bg.get(trend, "#ffffff")

        explanation = safe_text(row.get("Explanation", ""))
        explanation_html = explanation.replace("\n", "<br>")

        html += f"""
          <tr>
            <td style="border:1px solid #ddd; padding:8px; vertical-align:top;">{safe_text(row.get("Domain", ""))}</td>
            <td style="border:1px solid #ddd; padding:8px; vertical-align:top;">{safe_text(row.get("Strand", ""))}</td>
            <td style="
                border:1px solid #ddd;
                padding:8px;
                vertical-align:top;
                white-space:normal;
                word-wrap:break-word;
                overflow-wrap:break-word;
                line-height:1.4;
            ">{explanation_html}</td>
            <td style="border:1px solid #ddd; padding:8px; text-align:center; background:{initial_bg}; font-weight:bold;">{initial}</td>
            <td style="border:1px solid #ddd; padding:8px; text-align:center; background:{final_bg}; font-weight:bold;">{final}</td>
            <td style="border:1px solid #ddd; padding:8px; text-align:center; background:{trend_bg_color}; font-weight:bold;">{trend}</td>
          </tr>
        """

    html += """
        </tbody>
      </table>
    </div>
    """
    return html

def highlight_rating(val):
    color_map = {
        "HE": "#a8e6a1",   # green
        "E": "#d0f0fd",    # light blue
        "IN": "#fff3b0",   # yellow
        "DNMS": "#f8a5a5"  # red
    }
    return f"background-color: {color_map.get(val, '')}; color: black;"


def highlight_trend(val):
    if "Improved" in val:
        return "color: green; font-weight: 600;"
    elif "Dropped" in val:
        return "color: red; font-weight: 600;"
    elif "No change" in val:
        return "color: #555; font-weight: 500;"
    return ""

def render_grouped_comparison(df, key_prefix="cmp"):
    if df.empty:
        st.info("No comparison data available.")
        return

    domain_titles = {
        "A": "Planning and Preparation for Learning",
        "B": "Classroom Management",
        "C": "Delivery of Instruction",
        "D": "Monitoring, Assessment, and Follow-Up",
        "E": "Family and Community Outreach",
        "F": "Professional Responsibility",
    }

    domain_order = ["A", "B", "C", "D", "E", "F"]

    for domain in domain_order:
        domain_df = df[df["Domain"] == domain].copy()
        if domain_df.empty:
            continue

        display_df = domain_df[["Strand", "Initial", "Final", "Trend"]].copy()

        styled_df = (
            display_df.style
            .map(highlight_rating, subset=["Initial", "Final"])
            .map(highlight_trend, subset=["Trend"])
            .set_properties(subset=["Initial", "Final", "Trend"], **{
                "text-align": "center",
                "padding": "6px",
                "font-size": "13px"
            })
            .set_properties(subset=["Strand"], **{
                "padding": "6px",
                "font-size": "13px"
            })
        )

        expander_title = f"Domain {domain} — {domain_titles.get(domain, '')}"

        with st.expander(expander_title, expanded=(domain == "A")):
            st.dataframe(
                styled_df,
                use_container_width=True,
                hide_index=True
            )

def build_printable_comparison_html(teacher_name, teacher_email, appraiser, latest_initial, latest_final, display_df):
    initial_date = ""
    final_date = ""

    if latest_initial is not None and not latest_initial.empty:
        initial_date = safe_text(latest_initial.iloc[0].get("Timestamp", ""))

    if latest_final is not None and not latest_final.empty:
        final_date = safe_text(latest_final.iloc[0].get("Timestamp", ""))

    table_html = render_comparison_html(display_df)

    html = f"""
    <html>
    <head>
        <title>{teacher_name} - Initial vs Final Comparison</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                margin: 24px;
                color: #111;
            }}
            h1 {{
                font-size: 24px;
                margin-bottom: 8px;
            }}
            h2 {{
                font-size: 18px;
                margin-top: 0;
                margin-bottom: 20px;
                color: #444;
            }}
            .meta {{
                margin-bottom: 20px;
                line-height: 1.6;
                font-size: 14px;
            }}
            .meta strong {{
                display: inline-block;
                min-width: 140px;
            }}
            .print-btn {{
                margin-bottom: 20px;
            }}
            @media print {{
                .print-btn {{
                    display: none;
                }}
                body {{
                    margin: 10mm;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="print-btn">
            <button onclick="window.print()" style="padding:10px 16px; font-size:14px; cursor:pointer;">
                Print
            </button>
        </div>

        <h1>{teacher_name}</h1>
        <h2>Initial vs Final Self-Assessment Comparison</h2>

        <div class="meta">
            <div><strong>Email:</strong> {teacher_email}</div>
            <div><strong>Appraiser:</strong> {appraiser}</div>
            <div><strong>Initial Submitted:</strong> {initial_date or "-"}</div>
            <div><strong>Final Submitted:</strong> {final_date or "-"}</div>
        </div>

        {table_html}
    </body>
    </html>
    """
    return html

# =========================
# FINAL EVALUATION HELPERS
# =========================
def count_words(text):
    return len(re.findall(r"\b\S+\b", safe_text(text)))


def is_before_deadline(deadline_dt):
    return datetime.now() <= deadline_dt


def teacher_can_start_final_evaluation(email: str) -> bool:
    return user_has_submission(email, cycle="Final")


def final_eval_expected_headers():
    return [
        "Timestamp",
        "Last Edited On",
        "Teacher Email",
        "Teacher Name",
        "Appraiser",
        "Subject Area",
        "Student Survey Feedback",
        "Overall Reflection",
        "Teacher Submitted",
        "Teacher Submitted On",
        "Appraiser Started",
        "Appraiser Completed",
        "Appraiser Completed On",
        "A Rating",
        "B Rating",
        "C Rating",
        "D Rating",
        "E Rating",
        "F Rating",
        "Overall Rating",
        "Overall Comments",
        "Evaluator Sign Off",
        "Evaluator Sign Off Date",
        "Teacher Sign Off",
        "Teacher Sign Off Date",
    ]


@st.cache_resource
def ensure_final_eval_headers_once():
    exp = final_eval_expected_headers()
    current = with_backoff(FINAL_EVAL_WS.row_values, 1)

    if not current:
        with_backoff(FINAL_EVAL_WS.insert_row, exp, 1)
        return True

    if current != exp:
        st.warning(
            "The existing header row in **FinalEvaluation** does not match the expected structure. "
            "Submissions may misalign if the header was changed manually."
        )

    return True


@st.cache_data(ttl=180)
def load_final_eval_df():
    vals = with_backoff(FINAL_EVAL_WS.get_all_values)

    if not vals:
        return pd.DataFrame(columns=final_eval_expected_headers())

    header, rows = vals[0], vals[1:]
    df = pd.DataFrame(rows, columns=header) if rows else pd.DataFrame(columns=header)

    if "Teacher Email" in df.columns:
        df["Teacher Email"] = df["Teacher Email"].astype(str).str.strip().str.lower()

    if "Appraiser" in df.columns:
        df["Appraiser"] = df["Appraiser"].astype(str).str.strip().str.lower()

    return df


def get_teacher_final_eval_record(teacher_email: str):
    df = load_final_eval_df()

    if df.empty:
        return {}

    teacher_email = teacher_email.strip().lower()
    rows = df[df["Teacher Email"] == teacher_email]

    if rows.empty:
        return {}

    if "Timestamp" in rows.columns:
        rows = rows.sort_values("Timestamp", ascending=False)

    return dict(rows.iloc[0])


def save_final_eval_record(record: dict):
    headers = final_eval_expected_headers()
    df = load_final_eval_df()

    teacher_email = safe_text(record.get("Teacher Email", "")).strip().lower()
    row_values = [record.get(col, "") for col in headers]

    if not df.empty and "Teacher Email" in df.columns:
        matches = df[df["Teacher Email"].astype(str).str.strip().str.lower() == teacher_email]

        if not matches.empty:
            row_num = matches.index[0] + 2
            with_backoff(FINAL_EVAL_WS.update, f"A{row_num}:Y{row_num}", [row_values])
            load_final_eval_df.clear()
            return

    with_backoff(FINAL_EVAL_WS.append_row, row_values, value_input_option="USER_ENTERED")
    load_final_eval_df.clear()


def teacher_final_eval_completed(teacher_email: str) -> bool:
    rec = get_teacher_final_eval_record(teacher_email)
    return safe_text(rec.get("Teacher Submitted", "")).strip().lower() == "yes"


def appraiser_final_eval_completed(teacher_email: str) -> bool:
    rec = get_teacher_final_eval_record(teacher_email)
    return safe_text(rec.get("Appraiser Completed", "")).strip().lower() == "yes"


def evaluator_signed_off(teacher_email: str) -> bool:
    rec = get_teacher_final_eval_record(teacher_email)
    return safe_text(rec.get("Evaluator Sign Off", "")).strip().lower() == "yes"


def teacher_signed_off_final_eval(teacher_email: str) -> bool:
    rec = get_teacher_final_eval_record(teacher_email)
    return safe_text(rec.get("Teacher Sign Off", "")).strip().lower() == "yes"


def final_eval_domain_rows():
    return [
        ("A Rating", "A. Planning and Preparation for Learning"),
        ("B Rating", "B. Classroom Management"),
        ("C Rating", "C. Delivery of Instruction"),
        ("D Rating", "D. Monitoring, Assessment, and Follow-Up"),
        ("E Rating", "E. Family and Community Outreach"),
        ("F Rating", "F. Professional Responsibilities"),
    ]

def teacher_started_final_evaluation(teacher_email: str) -> bool:
    rec = get_teacher_final_eval_record(teacher_email)
    if not rec:
        return False

    return any([
        safe_text(rec.get("Teacher Submitted", "")).strip().lower() == "yes",
        safe_text(rec.get("Subject Area", "")).strip() != "",
        safe_text(rec.get("Student Survey Feedback", "")).strip() != "",
        safe_text(rec.get("Overall Reflection", "")).strip() != "",
    ])


def teacher_can_edit_final_self_assessment(teacher_email: str) -> bool:
    if not user_has_submission(teacher_email, cycle="Final"):
        return False

    if teacher_started_final_evaluation(teacher_email):
        return False

    return True

def domain_letter_from_strand(strand_code: str) -> str:
    return safe_text(strand_code).split()[0][:1]


def get_full_appraiser_name(appraiser_value: str) -> str:
    raw = safe_text(appraiser_value).strip()
    if not raw:
        return "Not Assigned"

    parts = [p.strip().lower() for p in raw.split(",") if p.strip()]
    if not parts:
        return raw

    matched_names = []

    for part in parts:
        match = users_df[
            users_df["Name"].astype(str).str.strip().str.lower().str.split().str[0] == part
        ]

        if not match.empty:
            matched_names.extend(match["Name"].astype(str).tolist())
        else:
            matched_names.append(part.title())

    return ", ".join(dict.fromkeys(matched_names))

def render_final_evaluation_review_panel(record: dict, heading: str = "Appraiser Review"):
    rating_colour_map = {
        "Highly Effective": "#d4edda",
        "Effective": "#d1ecf1",
        "Improvement Necessary": "#fff3cd",
        "Does Not Meet Standards": "#f8d7da",
    }

    text_colour_map = {
        "Highly Effective": "#155724",
        "Effective": "#0c5460",
        "Improvement Necessary": "#856404",
        "Does Not Meet Standards": "#721c24",
    }

    st.markdown(f"### {heading}")
    st.markdown("#### Ratings on Individual Rubrics")

    cols = st.columns(2)
    domain_rows = final_eval_domain_rows()

    for i, (col_name, label) in enumerate(domain_rows):
        rating_value = safe_text(record.get(col_name, ""))
        bg = rating_colour_map.get(rating_value, "#f4f4f4")
        fg = text_colour_map.get(rating_value, "#222")

        with cols[i % 2]:
            st.markdown(
                f"""
                <div style="
                    border: 1px solid #e6e6e6;
                    border-radius: 12px;
                    padding: 14px 16px;
                    margin-bottom: 12px;
                    background: #ffffff;
                    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
                ">
                    <div style="
                        font-size: 14px;
                        font-weight: 600;
                        color: #333;
                        margin-bottom: 10px;
                    ">
                        {label}
                    </div>
                    <div style="
                        display: inline-block;
                        padding: 8px 12px;
                        border-radius: 999px;
                        background: {bg};
                        color: {fg};
                        font-weight: 700;
                        font-size: 13px;
                    ">
                        {rating_value}
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )

    st.markdown("### Overall Rating")

    overall_rating = safe_text(record.get("Overall Rating", ""))
    overall_bg = rating_colour_map.get(overall_rating, "#f4f4f4")
    overall_fg = text_colour_map.get(overall_rating, "#222")

    st.markdown(
        f"""
        <div style="
            border: 2px solid #dcdcdc;
            border-radius: 14px;
            padding: 18px;
            margin-top: 8px;
            margin-bottom: 14px;
            background: #fafafa;
            box-shadow: 0 1px 6px rgba(0,0,0,0.05);
        ">
            <div style="
                font-size: 15px;
                font-weight: 600;
                color: #333;
                margin-bottom: 12px;
            ">
                Final Overall Rating
            </div>
            <div style="
                display: inline-block;
                padding: 10px 16px;
                border-radius: 999px;
                background: {overall_bg};
                color: {overall_fg};
                font-weight: 700;
                font-size: 15px;
            ">
                {overall_rating}
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown("### Appraiser Comments")
    st.markdown(
        f"""
        <div style="
            border: 1px solid #e6e6e6;
            border-radius: 12px;
            padding: 16px;
            background: #ffffff;
            box-shadow: 0 1px 4px rgba(0,0,0,0.06);
            line-height: 1.6;
            color: #333;
            margin-bottom: 12px;
        ">
            {safe_text(record.get("Overall Comments", "")).replace("\n", "<br>")}
        </div>
        """,
        unsafe_allow_html=True
    )
    
# =========================
# UI CONFIG (must be first)
# =========================
st.set_page_config(page_title="OIS Teacher Appraisal", layout="wide")

inject_css()

# Try to import HttpError; fall back gracefully if googleapiclient isn't present
try:
    from googleapiclient.errors import HttpError  # type: ignore
except Exception:  # pragma: no cover
    class HttpError(Exception):
        pass

# =========================
# RERUN helper (Streamlit API changed)
# =========================
def _rerun():
    try:
        st.rerun()  # Streamlit >=1.32
    except AttributeError:
        st.experimental_rerun()  # Older versions


# =========================
# CONFIG
# =========================
SPREADSHEET_ID = "1kqcfnMx4KhqQvFljsTwSOcmuEHnkLAdwp_pUJypOjpY"
SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]
ENABLE_REFLECTIONS = True  # set False to hide reflection boxes
CURRENT_ASSESSMENT_CYCLE = "Final"   # "Initial" or "Final"

FINAL_EVAL_SHEET_NAME = "FinalEvaluation"

FINAL_EVAL_TEACHER_DEADLINE = datetime(2026, 4, 30, 23, 59, 59)
FINAL_EVAL_APPRAISER_DEADLINE = datetime(2026, 5, 20, 23, 59, 59)

FINAL_EVAL_MAX_WORDS_SURVEY = 150
FINAL_EVAL_MAX_WORDS_REFLECTION = 150
FINAL_EVAL_MAX_WORDS_COMMENTS = 150

FINAL_EVAL_RATINGS = [
    "Highly Effective",
    "Effective",
    "Improvement Necessary",
    "Does Not Meet Standards",
]

SUBJECT_AREA_OPTIONS = [
    "English",
    "Mathematics",
    "Science",
    "Individuals and Societies",
    "Languages",
    "Design",
    "Physical and Health Education",
    "Visual Arts",
    "Music",
    "Theatre",
    "Computer Science",
    "SSP",
    "Other",
]

# Optional: list of admin emails (lowercase) in .streamlit/secrets.toml
ADMINS_FROM_SECRETS = set([e.strip().lower() for e in st.secrets.get("admins", [])])

IST_OFFSET_HOURS = 5
IST_OFFSET_MINUTES = 30

def now_ist():
    return datetime.utcnow() + pd.Timedelta(hours=IST_OFFSET_HOURS, minutes=IST_OFFSET_MINUTES)

def now_ist_str():
    return now_ist().strftime("%Y-%m-%d %H:%M:%S")

def fmt_ist(dt_value):
    txt = safe_text(dt_value)
    return txt if txt else "-"

# =========================
# DOMAINS & SUB-STRANDS (exact from rubric)
# =========================
DOMAINS = {
    "A: Planning and Preparation for Learning": [
        ("A1", "Expertise"),
        ("A2", "Goals"),
        ("A3", "Units"),
        ("A4", "Assessments"),
        ("A5", "Anticipation"),
        ("A6", "Lessons"),
        ("A7", "Materials"),
        ("A8", "Differentiation"),
        ("A9", "Environment"),
    ],
    "B: Classroom Management": [
        ("B1", "Expectations"),
        ("B2", "Relationships"),
        ("B3", "Social Emotional"),
        ("B4", "Routines"),
        ("B5", "Responsibility"),
        ("B6", "Repertoire"),
        ("B7", "Prevention"),
        ("B8", "Incentives"),
    ],
    "C: Delivery of Instruction": [
        ("C1", "Expectations"),
        ("C2", "Mindset"),
        ("C3", "Framing"),
        ("C4", "Connections"),
        ("C5", "Clarity"),
        ("C6", "Repertoire"),
        ("C7", "Engagement"),
        ("C8", "Differentiation"),
        ("C9", "Nimbleness"),
    ],
    "D: Monitoring, Assessment, and Follow-Up": [
        ("D1", "Criteria"),
        ("D2", "Diagnosis"),
        ("D3", "Goals"),
        ("D4", "Feedback"),
        ("D5", "Recognition"),
        ("D6", "Analysis"),
        ("D7", "Tenacity"),
        ("D8", "Support"),
        ("D9", "Reflection"),
    ],
    "E: Family and Community Outreach": [
        ("E1", "Respect"),
        ("E2", "Belief"),
        ("E3", "Expectations"),
        ("E4", "Communication"),
        ("E5", "Involving"),
        ("E6", "Responsiveness"),
        ("E7", "Reporting"),
        ("E8", "Outreach"),
        ("E9", "Resources"),
    ],
    "F: Professional Responsibility": [
        ("F1", "Language"),
        ("F2", "Reliability"),
        ("F3", "Professionalism"),
        ("F4", "Judgement"),
        ("F5", "Teamwork"),
        ("F6", "Leadership"),
        ("F7", "Openness"),
        ("F8", "Collaboration"),
        ("F9", "Growth"),
    ],
}

# Rating scale (exact rubric wording)
RATINGS = [
    "Highly Effective",
    "Effective",
    "Improvement Necessary",
    "Does Not Meet Standards",
]

# =========================
# Small retry/backoff for Sheets calls (handles 429/5xx)
# =========================
def with_backoff(fn, *args, **kwargs):
    """Retry gspread/api calls briefly on 429/5xx."""
    max_attempts = 5
    delay = 0.6  # seconds
    last_exc = None
    for _ in range(max_attempts):
        try:
            return fn(*args, **kwargs)
        except HttpError as e:  # googleapiclient
            status = getattr(e, "status_code", None)
            if status in (429, 500, 502, 503, 504):
                time.sleep(delay); delay *= 2; last_exc = e; continue
            raise
        except gspread.exceptions.APIError as e:  # gspread-wrapped
            msg = str(e).lower()
            if any(code in msg for code in ["429", "500", "502", "503", "504"]):
                time.sleep(delay); delay *= 2; last_exc = e; continue
            raise
        except Exception as e:
            time.sleep(delay); delay *= 2; last_exc = e; continue
    if last_exc:
        raise last_exc
    return fn(*args, **kwargs)

# =========================
# ONE-TIME SHEETS CONNECTION (cached)
# =========================
# @st.cache_resource
# =========================
# Google Sheet Connections
# =========================
@st.cache_resource
def get_worksheets():
    client = gspread.authorize(
        Credentials.from_service_account_info(st.secrets["google"], scopes=SCOPES)
    )
    sh = client.open_by_key(SPREADSHEET_ID)

    resp_ws = sh.worksheet("Responses")
    users_ws = sh.worksheet("Users")

    try:
        drafts_ws = sh.worksheet("Drafts")
    except gspread.exceptions.WorksheetNotFound:
        drafts_ws = sh.add_worksheet(title="Drafts", rows="1000", cols="100")
        drafts_ws.update([["Email"]])

    try:
        final_eval_ws = sh.worksheet(FINAL_EVAL_SHEET_NAME)
    except gspread.exceptions.WorksheetNotFound:
        final_eval_ws = sh.add_worksheet(title=FINAL_EVAL_SHEET_NAME, rows="1000", cols="50")

    return resp_ws, users_ws, drafts_ws, final_eval_ws
    
RESP_WS, USERS_WS, DRAFTS_WS, FINAL_EVAL_WS = get_worksheets()

# =========================
# DRAFT HELPERS
# =========================
def save_draft(email, form_data):
    """Update or append a draft for this teacher only."""
    try:
        # Get all drafts (lightweight, header + values)
        all_drafts = DRAFTS_WS.get_all_records()
        emails = [row["Email"] for row in all_drafts]

        row_data = [email] + [form_data.get(f, "") for f in form_data.keys()]

        if email in emails:
            # Update existing row (Google Sheets is 1-indexed and has a header row)
            row_num = emails.index(email) + 2  
            DRAFTS_WS.update(f"A{row_num}", [row_data])
        else:
            # Append new row
            if not all_drafts:  
                # If sheet is empty except header, add header first
                headers = ["Email"] + list(form_data.keys())
                DRAFTS_WS.append_row(headers, value_input_option="USER_ENTERED")
            DRAFTS_WS.append_row(row_data, value_input_option="USER_ENTERED")

        return True
    except Exception as e:
        st.error(f"⚠️ Could not save draft: {e}")
        return False


def load_draft(email):
    """Load teacher's draft if exists."""
    try:
        all_drafts = pd.DataFrame(DRAFTS_WS.get_all_records())
        user_draft = all_drafts[all_drafts["Email"] == email]
        if not user_draft.empty:
            return dict(user_draft.iloc[0])
    except Exception:
        return {}
    return {}

# =========================
# HEADER MANAGEMENT (safe, non-destructive)
# =========================
def expected_headers():
    headers = ["Timestamp", "Email", "Name", "Appraiser", "Assessment Cycle"]
    for domain, items in DOMAINS.items():
        for code, label in items:
            headers.append(f"{code} {label}")
        if ENABLE_REFLECTIONS:
            headers.append(f"{domain} Reflection")
    headers.append("Last Edited On")
    return headers

@st.cache_resource
def ensure_headers_once():
    exp = expected_headers()
    current = with_backoff(RESP_WS.row_values, 1)
    if not current:
        with_backoff(RESP_WS.insert_row, exp, 1)
        return True
    if current != exp:
        st.warning(
            "The existing header row in **Responses** does not match the current rubric. "
            "Submissions will still append, but columns may be misaligned if the rubric changed. "
            "To update safely, export data, fix headers offline, and re-import."
        )
    return True

ensure_headers_once()
ensure_final_eval_headers_once()

# =========================
# USERS: read ONCE per server process (auto‑detect headers)
# =========================
def _pick_col(candidates: list[str], cols: list[str]):
    norm_map = {c.strip().lower(): c for c in cols}
    for want in candidates:
        key = want.strip().lower()
        if key in norm_map: return norm_map[key]
    for c in cols:
        cl = c.strip().lower()
        if any(w in cl for w in candidates): return c
    return None

@st.cache_resource
def load_users_once_df():
    """
    Load Users sheet once and normalise key columns, including Campus.

    Expected logical columns (case-insensitive / fuzzy matched):
      - Email
      - Name
      - Appraiser
      - Role
      - Password
      - Campus  (NEW – optional; if missing or blank → treated as single-campus setup)
    """
    records = with_backoff(USERS_WS.get_all_records)
    if not records:
        return pd.DataFrame(columns=["Email", "Name", "Appraiser", "Role", "Password", "Campus"])

    df = pd.DataFrame(records)
    if df.empty:
        return pd.DataFrame(columns=["Email", "Name", "Appraiser", "Role", "Password", "Campus"])

    cols = list(df.columns)

    email_header     = _pick_col(["email", "school email", "work email", "ois email", "e-mail"], cols)
    name_header      = _pick_col(["name", "full name", "teacher name", "staff name"], cols)
    appraiser_header = _pick_col(["appraiser", "line manager", "manager", "appraiser name", "supervisor"], cols)
    role_header      = _pick_col(["role", "access", "admin"], cols)
    password_header  = _pick_col(["password", "pwd", "pass"], cols)
    campus_header    = _pick_col(["campus"], cols)  # NEW

    out = pd.DataFrame()

    # Core columns
    out["Email"] = (
        df[email_header].astype(str).str.strip().str.lower()
        if email_header else ""
    )
    out["Name"] = (
        df[name_header].astype(str).str.strip()
        if name_header else ""
    )
    out["Appraiser"] = (
        df[appraiser_header].astype(str).str.strip().replace({"": "Not Assigned"})
        if appraiser_header else "Not Assigned"
    )
    out["Role"] = (
        df[role_header].astype(str).str.strip().str.lower()
        if role_header else ""
    )
    out["Password"] = (
        df[password_header].astype(str).str.strip()
        if password_header else ""
    )

    # NEW: Campus (e.g. "JVLR" / "OGC")
    out["Campus"] = (
        df[campus_header].astype(str).str.strip()
        if campus_header else ""
    )

    return out


users_df = load_users_once_df()


# =========================
# RESPONSES cache (for 'My submission' and Admin)
# =========================
@st.cache_data(ttl=180)  # slightly longer to reduce bursts
def load_responses_df():
    vals = with_backoff(RESP_WS.get_all_values)
    if not vals:
        return pd.DataFrame()

    header, rows = vals[0], vals[1:]
    df = pd.DataFrame(rows, columns=header) if rows else pd.DataFrame(columns=header)

    if "Email" in df.columns:
        df["Email"] = df["Email"].astype(str).str.lower()

    # Backfill older records that were submitted before Assessment Cycle existed
    if "Assessment Cycle" not in df.columns:
        df["Assessment Cycle"] = "Initial"
    else:
        df["Assessment Cycle"] = df["Assessment Cycle"].replace("", "Initial")

    return df

def user_has_submission(email: str, cycle: str | None = None) -> bool:
    if not email:
        return False

    df = load_responses_df()
    if df.empty or "Email" not in df.columns:
        return False

    filtered = df[df["Email"] == email.strip().lower()]

    if cycle is not None and "Assessment Cycle" in df.columns:
        filtered = filtered[filtered["Assessment Cycle"] == cycle]

    return not filtered.empty

# =========================
# Authentication & Roles
# =========================
def authenticate_user(email, password):
    email = email.strip().lower()

    # Look up in Users sheet
    user_row = users_df[users_df["Email"].str.lower() == email]
    if user_row.empty:
        return None, None  # not found

    role = user_row.iloc[0]["Role"].strip().lower()

    # Admin check
    if role == "admin":
        return ("admin", user_row.iloc[0]) if password == "OIS2025" else (None, None)

    # Superadmin check
    if role == "sadmin":
        return ("sadmin", user_row.iloc[0]) if password == "SOIS2025" else (None, None)

    # Teacher check — validate against Password column
    if role == "user":
        stored_pw = str(user_row.iloc[0].get("Password", "")).strip()
        entered_pw = str(password).strip()

        if stored_pw and entered_pw and stored_pw == entered_pw:
            return "user", user_row.iloc[0]
        else:
            st.warning(f"Debug → Entered: '{entered_pw}', Stored: '{stored_pw}'")
            return None, None


# ============ TOP OF SIDEBAR: USER + CAMPUS ============
user_name = st.session_state.get("auth_name", "")
campus_label = st.session_state.get("auth_campus", "")

with st.sidebar:
    st.markdown("### 👤 Logged in as")
    if user_name:
        st.markdown(f"**{user_name}**")
    if campus_label:
        st.markdown(f"🏫 **{campus_label} Campus**")
    
# =========================
# AUTH: Account + Logout (from Google login in app.py)
# =========================
if "auth_email" not in st.session_state or not st.session_state.auth_email:
    st.info("Please log in first.")
    st.stop()
    
if st.sidebar.button("🚪 **LOGOUT**", type="primary", use_container_width=True):
    # Clear all login-related session keys
    for key in ["token", "auth_email", "auth_name", "auth_role", "auth_campus", "submitted"]:
        if key in st.session_state:
            del st.session_state[key]


    st.cache_data.clear()
    st.cache_resource.clear()

    # Force redirect to app.py (login)
    st.switch_page("app.py")

     

# =========================
# Sidebar: Live progress (teachers only)
# =========================
total_items = sum(len(v) for v in DOMAINS.values())

def current_progress_from_session() -> int:
    count = 0
    for _, items in DOMAINS.items():
        for code, label in items:
            if st.session_state.get(f"{code}-{label}"):
                count += 1
    return count

if st.session_state.get("auth_role") == "user":
    with st.sidebar.expander("📊 Progress", expanded=True):
        done = current_progress_from_session()
        st.progress(done / total_items if total_items else 0.0)
        st.caption(f"{done}/{total_items} sub-strands completed")

    # Show initial reference panel in sidebar during Final cycle
    if CURRENT_ASSESSMENT_CYCLE == "Final":
        _ref_initial, _, _ = build_teacher_initial_final(st.session_state.auth_email)
        if _ref_initial is not None and not _ref_initial.empty:
            _ref_record = _ref_initial.iloc[0].to_dict()
            _short_map = {
                "Highly Effective": "HE", "Effective": "E",
                "Improvement Necessary": "IN", "Does Not Meet Standards": "DNMS"
            }
            _colour_map = {"HE": "🟩", "E": "🟦", "IN": "🟨", "DNMS": "🟥"}
            st.sidebar.markdown("---")
            st.sidebar.markdown("**📘 Initial reference · Sep 2025**")
            for _d, _items in DOMAINS.items():
                _dletter = _d.split(":")[0]
                with st.sidebar.expander(f"Domain {_dletter}", expanded=False):
                    for _code, _label in _items:
                        _strand = f"{_code} {_label}"
                        _val = _ref_record.get(_strand, "")
                        _short = _short_map.get(_val, _val)
                        _col = _colour_map.get(_short, "⬜")
                        st.markdown(f"{_col} **{_code}** {_label} — `{_short}`")
            st.sidebar.markdown(
                '<a href="https://drive.google.com/file/d/1GrDAkk8zev6pr4AmmKA6YyTzeUdZ8dZC/view?usp=sharing" ' +
                'target="_blank" style="font-size:12px;color:#6d28d9;font-weight:600;">📄 View Teacher Growth Rubric</a>',
                unsafe_allow_html=True
            )
        
# Main Nav
st.markdown("""
<div style="display:flex;align-items:center;gap:12px;margin-bottom:4px">
  <span style="font-size:26px">🌟</span>
  <div>
    <div style="font-size:20px;font-weight:700;color:#111827;line-height:1.2">OIS Teacher Appraisal 2025–26</div>
    <div style="font-size:13px;color:#6b7280">Oberoi International School · JVLR Campus</div>
  </div>
</div>
<hr style="border:none;border-top:1px solid #e5e7eb;margin:10px 0 18px"/>
""", unsafe_allow_html=True)

if not st.session_state.auth_email:
    st.info("Please log in from the sidebar to continue.")
    st.stop()

already_submitted = user_has_submission(
    st.session_state.auth_email,
    cycle=CURRENT_ASSESSMENT_CYCLE
)

# Look up my role (and campus, if configured) from the Users table
me_row = users_df[users_df["Email"] == st.session_state.auth_email]
if me_row.empty:
    role = "user"
    campus = ""
else:
    role = str(me_row.iloc[0].get("Role", "user")).lower().strip()
    campus = str(me_row.iloc[0].get("Campus", "")).strip()

     
# Mirror into session (login also sets this)
st.session_state.auth_role = role
st.session_state.auth_campus = campus

i_am_admin = (role == "admin")
i_am_sadmin = (role == "sadmin")

# Decide which tabs to show
if i_am_sadmin:
    nav_options = ["Super Admin"]
elif i_am_admin:
    nav_options = ["Admin"]
else:
    teacher_has_final_self = teacher_can_start_final_evaluation(st.session_state.auth_email)

    if already_submitted:
        nav_options = ["My Submission"]
    else:
        nav_options = ["Self-Assessment (Initial & Final)", "My Submission"]

    nav_options.append("Final Evaluation")

    if not teacher_has_final_self:
        st.sidebar.caption("Final Evaluation unlocks after Final self-assessment submission.")

# This defines `tab`
tab = st.sidebar.radio("Menu", nav_options, index=0)
admin_view_mode = None

if tab == "Admin" and i_am_admin:
    admin_view_mode = st.sidebar.selectbox(
        "Jump to",
        ["Summary of Teachers", "View Teacher Self-Assessment", "Self-Assessment Grid"],
        index=0
    )

sadmin_view_mode = None

if tab == "Super Admin" and i_am_sadmin:
    sadmin_view_mode = st.sidebar.selectbox(
        "Jump to",
        ["Summary of Teachers", "View Teacher Self-Assessment", "Self-Assessment Grid"],
        index=0
    )

# =========================
# Page: Self-Assessment (teachers who haven't submitted yet)
# =========================
from descriptors import DESCRIPTORS  # 👈 make sure descriptors.py is in same folder

if tab == "Self-Assessment (Initial & Final)":
    if already_submitted and not i_am_admin:
        # Auto-redirect teachers with submissions to My Submission
        st.success("✅ You’ve already submitted your self-assessment. Redirecting to your submission...")
        tab = "My Submission"
    else:
        # Welcome + Appraiser info
        me = users_df[users_df["Email"] == st.session_state.auth_email].iloc[0] if not users_df.empty else {}
        appraiser = me.get("Appraiser", "Not Assigned") if isinstance(me, pd.Series) else "Not Assigned"
        st.sidebar.info(f"Your appraiser: **{appraiser}**")

        # Step tracker
        if CURRENT_ASSESSMENT_CYCLE == "Final":
            render_step_tracker([
                ("Initial", "Sep 2025 — submitted ✓", "done"),
                ("Final self-assessment", "Apr 2026 — in progress", "active"),
                ("Final evaluation", "Unlocks after step 2", "locked"),
                ("Sign-off", "After your meeting", "locked"),
            ])
            render_guidance(
                "How to complete this",
                "Your <strong>Initial self-assessment (Sep 2025)</strong> is in the sidebar for reference. "
                "Rate yourself on each strand independently — your appraiser cannot see this until you submit. "
                "Open any strand to view the full descriptors before choosing your rating."
            )
        else:
            render_step_tracker([
                ("Initial self-assessment", "Sep 2025 — in progress", "active"),
                ("Final self-assessment", "Apr 2026", "locked"),
                ("Final evaluation", "After final submission", "locked"),
                ("Sign-off", "After your meeting", "locked"),
            ])
            render_guidance(
                "How to complete this",
                "Rate yourself honestly on each strand using the Kim Marshall rubric descriptors. "
                "You can save a draft and return — only submit when all strands are rated."
            )

        # 🔹 Load draft if exists
        draft_data = load_draft(st.session_state.auth_email) or {}
        
        # 🔹 ALWAYS load initial/final (not inside draft condition)
        latest_initial, latest_final, comparison_df = build_teacher_initial_final(
            st.session_state.auth_email
        )
        
        if draft_data:
            st.info("💾 A saved draft was found and preloaded. You can continue where you left off.")
        
        
        
            st.divider()
            st.markdown("### Final Self-Assessment - Apr 2026")
            st.caption("Please complete your final self-assessment independently. Use your initial self-assessment on the sidebar as a reference.")
            st.info("""
            **Please complete the Final Self-Assessment in this order:**
            
            **Step 1:** Review your Initial Self-Assessment from Sep 2025 above and in the sidebar reference.  
            **Step 2:** Reflect on your practice in each domain before selecting ratings.  
            **Step 3:** Complete all strands across Domains A–F.  
            **Step 4:** Review your selections carefully before submitting.  
            """)

        # Selections (direct widgets so sidebar progress updates live)
        selections = {}
        reflections = {}

        for domain, items in DOMAINS.items():
            with st.expander(domain, expanded=False):
                for code, label in items:
                    strand_key = f"{code} {label}"
                    key = f"{code}-{label}"
                    saved_value = draft_data.get(strand_key, "")

                    # Show initial rating inline (Final cycle only)
                    if CURRENT_ASSESSMENT_CYCLE == "Final" and latest_initial is not None and not latest_initial.empty:
                        _init_v = safe_text(latest_initial.iloc[0].get(strand_key, ""))
                        _init_s = rating_short(_init_v)
                        _pill_cls = {"HE": "rp-he", "E": "rp-e", "IN": "rp-in", "DNMS": "rp-dn"}.get(_init_s, "")
                        _pill = f'<span class="rating-pill {_pill_cls}" style="font-size:10px">{_init_s}</span>' if _init_s else ""
                        st.markdown(
                            f'<div style="font-size:12px;color:#6b7280;margin-bottom:2px">Initial: {_pill}</div>',
                            unsafe_allow_html=True
                        )

                    # Radio for selecting rating
                    selections[strand_key] = st.radio(
                        f"{strand_key}",
                        RATINGS,
                        index=RATINGS.index(saved_value) if saved_value in RATINGS else None,
                        key=key,
                        horizontal=True,
                    ) or ""

                    # 🔹 Show descriptors (auto-expand if no saved choice yet)
                    if strand_key in DESCRIPTORS:
                        expand_default = saved_value == ""  # open first time, collapse later
                        with st.expander("📖 See descriptors for this strand", expanded=expand_default):
                            st.markdown(f"""
                            **Highly Effective (HE):** {DESCRIPTORS[strand_key]['HE']}  

                            **Effective (E):** {DESCRIPTORS[strand_key]['E']}  

                            **Improvement Necessary (IN):** {DESCRIPTORS[strand_key]['IN']}  

                            **Does Not Meet Standards (DNMS):** {DESCRIPTORS[strand_key]['DNMS']}  
                            """)

                # Reflection box per domain (if enabled)
                if ENABLE_REFLECTIONS:
                    saved_refl = draft_data.get(f"Reflection-{domain}", "")
                    reflections[domain] = st.text_area(
                        f"{domain} Reflection (optional)",
                        key=f"refl-{domain}",
                        placeholder="Notes / evidence / next steps (optional)",
                        value=saved_refl,
                    )

        # Submit button + progress
        selected_count = sum(1 for v in selections.values() if v)
        col1, col2 = st.columns([1, 3])
        with col1:
            submit = st.button(
                "✅ Submit",
                disabled=(selected_count < total_items) or st.session_state.get("submitted", False)
            )

            # Sidebar: Save Draft
            with st.sidebar:
                if st.button("💾 Save Draft", use_container_width=True):
                    draft_payload = {}
                    for domain, items in DOMAINS.items():
                        for code, label in items:
                            draft_payload[f"{code} {label}"] = selections[f"{code} {label}"]
                        if ENABLE_REFLECTIONS:
                            draft_payload[f"Reflection-{domain}"] = reflections.get(domain, "")
                    save_draft(st.session_state.auth_email, draft_payload)
                    st.success("✅ Draft saved!")



        # Handle Submit
        if submit:
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
            row = [
                now_str,
                st.session_state.auth_email,
                st.session_state.auth_name,
                appraiser,
                CURRENT_ASSESSMENT_CYCLE,
            ]
        
            for domain, items in DOMAINS.items():
                for code, label in items:
                    row.append(selections[f"{code} {label}"])
                if ENABLE_REFLECTIONS:
                    row.append(reflections.get(domain, ""))
        
            row.append(now_str)  # Last Edited On
        
            try:
                with_backoff(RESP_WS.append_row, row, value_input_option="USER_ENTERED")
                load_responses_df.clear()
                st.session_state.submitted = True
                st.success("🎉 Submitted. Thank you! See **My Submission** to review your responses.")
            except Exception as e:
                st.error("⚠️ Could not submit right now. Please try again shortly.")
                st.caption(f"Debug info: {e}")

# =========================
# Page: My Submission (teachers see their data here)
# =========================
if tab == "My Submission":
    latest_initial, latest_final, comparison_df = build_teacher_initial_final(
        st.session_state.auth_email
    )

    if latest_initial is None and latest_final is None:
        st.info("No submission found yet. Complete your self-assessment to see your results here.")
    else:
        has_init = latest_initial is not None and not latest_initial.empty
        has_fin = latest_final is not None and not latest_final.empty
        fe_sub = teacher_final_eval_completed(st.session_state.auth_email)
        t_signed = teacher_signed_off_final_eval(st.session_state.auth_email)
        a_signed = evaluator_signed_off(st.session_state.auth_email)

        render_step_tracker([
            ("Initial", "Sep 2025", "done" if has_init else "locked"),
            ("Final self-assessment", "Apr 2026", "done" if has_fin else ("active" if has_init else "locked")),
            ("Final evaluation", "Your section", "done" if fe_sub else ("active" if has_fin else "locked")),
            ("Sign-off", "After meeting", "done" if (t_signed and a_signed) else "locked"),
        ])

        st.markdown('<div class="sec-heading">Initial vs Final Comparison</div>', unsafe_allow_html=True)

        if has_init and has_fin:
            init_date = safe_text(latest_initial.iloc[0].get("Timestamp", ""))
            fin_date = safe_text(latest_final.iloc[0].get("Timestamp", ""))
            st.markdown(
                f'<div style="font-size:12px;color:#6b7280;margin-bottom:10px">Initial: <strong>{init_date}</strong> &nbsp;·&nbsp; Final: <strong>{fin_date}</strong></div>',
                unsafe_allow_html=True
            )
            if not comparison_df.empty:
                display_df = comparison_df[["Domain", "Strand", "Explanation", "Initial", "Final", "Trend"]].copy()
                render_comparison_table_new(display_df)
        elif has_init and not has_fin:
            st.info("Your final self-assessment is not yet submitted. Once submitted, the comparison will appear here.")
            st.markdown("**Initial self-assessment:**")
            init_display = latest_initial.copy().replace({
                "Highly Effective": "HE", "Effective": "E",
                "Improvement Necessary": "IN", "Does Not Meet Standards": "DNMS"
            })
            st.dataframe(init_display.style.map(highlight_ratings, subset=init_display.columns[5:]), use_container_width=True)
        else:
            st.info("No initial submission available.")

  
# =========================
# Page: Final Evaluation (Teacher)
# =========================
if tab == "Final Evaluation" and role == "user":
    st.subheader("Final Evaluation")

    teacher_email = st.session_state.auth_email.strip().lower()
    teacher_name = st.session_state.auth_name

    me = users_df[users_df["Email"] == teacher_email].iloc[0] if not users_df.empty else {}
    appraiser_raw = me.get("Appraiser", "Not Assigned") if isinstance(me, pd.Series) else "Not Assigned"
    appraiser = get_full_appraiser_name(appraiser_raw)

    if not teacher_can_start_final_evaluation(teacher_email):
        st.warning("You must first submit your Final self-assessment before this section becomes active.")
        st.stop()

    record = get_teacher_final_eval_record(teacher_email)

    teacher_locked = (
        not is_before_deadline(FINAL_EVAL_TEACHER_DEADLINE)
        or teacher_final_eval_completed(teacher_email)
    )

    st.info(f"Appraiser: {appraiser}")
    st.caption(f"Teacher deadline (IST): {FINAL_EVAL_TEACHER_DEADLINE.strftime('%d %b %Y, %I:%M %p')}")

    subject_existing = safe_text(record.get("Subject Area", ""))
    survey_existing = safe_text(record.get("Student Survey Feedback", ""))
    reflection_existing = safe_text(record.get("Overall Reflection", ""))

    subject_index = SUBJECT_AREA_OPTIONS.index(subject_existing) if subject_existing in SUBJECT_AREA_OPTIONS else 0

    subject_area = st.selectbox(
        "Subject Area",
        SUBJECT_AREA_OPTIONS,
        index=subject_index,
        disabled=teacher_locked,
        key="fe_subject_area"
    )

    student_survey_feedback = st.text_area(
        "Student Survey Feedback (150 words or less)",
        value=survey_existing,
        height=180,
        disabled=teacher_locked,
        key="fe_student_survey"
    )
    survey_wc = count_words(student_survey_feedback)
    st.caption(f"Word count: {survey_wc}/{FINAL_EVAL_MAX_WORDS_SURVEY}")

    overall_reflection = st.text_area(
        "Overall Reflection on the school year (150 words or less)",
        value=reflection_existing,
        height=180,
        disabled=teacher_locked,
        key="fe_overall_reflection"
    )
    reflection_wc = count_words(overall_reflection)
    st.caption(f"Word count: {reflection_wc}/{FINAL_EVAL_MAX_WORDS_REFLECTION}")

    too_many_words = (
        survey_wc > FINAL_EVAL_MAX_WORDS_SURVEY
        or reflection_wc > FINAL_EVAL_MAX_WORDS_REFLECTION
    )

    if survey_wc > FINAL_EVAL_MAX_WORDS_SURVEY:
        st.error(f"Student Survey Feedback is over the {FINAL_EVAL_MAX_WORDS_SURVEY}-word limit.")

    if reflection_wc > FINAL_EVAL_MAX_WORDS_REFLECTION:
        st.error(f"Overall Reflection is over the {FINAL_EVAL_MAX_WORDS_REFLECTION}-word limit.")

    if too_many_words:
        st.warning("Buttons are disabled until both sections are within the word limit.")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("💾 Save Teacher Section", disabled=teacher_locked or too_many_words):
            now_str = now_ist_str()

            updated = {
                "Timestamp": safe_text(record.get("Timestamp", now_str)) or now_str,
                "Last Edited On": now_str,
                "Teacher Email": teacher_email,
                "Teacher Name": teacher_name,
                "Appraiser": appraiser,
                "Subject Area": subject_area,
                "Student Survey Feedback": student_survey_feedback,
                "Overall Reflection": overall_reflection,
                "Teacher Submitted": safe_text(record.get("Teacher Submitted", "")),
                "Teacher Submitted On": safe_text(record.get("Teacher Submitted On", "")),
                "Appraiser Started": safe_text(record.get("Appraiser Started", "")),
                "Appraiser Completed": safe_text(record.get("Appraiser Completed", "")),
                "Appraiser Completed On": safe_text(record.get("Appraiser Completed On", "")),
                "A Rating": safe_text(record.get("A Rating", "")),
                "B Rating": safe_text(record.get("B Rating", "")),
                "C Rating": safe_text(record.get("C Rating", "")),
                "D Rating": safe_text(record.get("D Rating", "")),
                "E Rating": safe_text(record.get("E Rating", "")),
                "F Rating": safe_text(record.get("F Rating", "")),
                "Overall Rating": safe_text(record.get("Overall Rating", "")),
                "Overall Comments": safe_text(record.get("Overall Comments", "")),
                "Evaluator Sign Off": safe_text(record.get("Evaluator Sign Off", "")),
                "Evaluator Sign Off Date": safe_text(record.get("Evaluator Sign Off Date", "")),
                "Teacher Sign Off": safe_text(record.get("Teacher Sign Off", "")),
                "Teacher Sign Off Date": safe_text(record.get("Teacher Sign Off Date", "")),
            }

            save_final_eval_record(updated)
            st.success("Teacher section saved.")
            _rerun()

    with col2:
        if st.button("✅ Submit Teacher Section", disabled=teacher_locked or too_many_words):
            now_str = now_ist_str()

            updated = {
                "Timestamp": safe_text(record.get("Timestamp", now_str)) or now_str,
                "Last Edited On": now_str,
                "Teacher Email": teacher_email,
                "Teacher Name": teacher_name,
                "Appraiser": appraiser,
                "Subject Area": subject_area,
                "Student Survey Feedback": student_survey_feedback,
                "Overall Reflection": overall_reflection,
                "Teacher Submitted": "Yes",
                "Teacher Submitted On": now_str,
                "Appraiser Started": safe_text(record.get("Appraiser Started", "")),
                "Appraiser Completed": safe_text(record.get("Appraiser Completed", "")),
                "Appraiser Completed On": safe_text(record.get("Appraiser Completed On", "")),
                "A Rating": safe_text(record.get("A Rating", "")),
                "B Rating": safe_text(record.get("B Rating", "")),
                "C Rating": safe_text(record.get("C Rating", "")),
                "D Rating": safe_text(record.get("D Rating", "")),
                "E Rating": safe_text(record.get("E Rating", "")),
                "F Rating": safe_text(record.get("F Rating", "")),
                "Overall Rating": safe_text(record.get("Overall Rating", "")),
                "Overall Comments": safe_text(record.get("Overall Comments", "")),
                "Evaluator Sign Off": safe_text(record.get("Evaluator Sign Off", "")),
                "Evaluator Sign Off Date": safe_text(record.get("Evaluator Sign Off Date", "")),
                "Teacher Sign Off": safe_text(record.get("Teacher Sign Off", "")),
                "Teacher Sign Off Date": safe_text(record.get("Teacher Sign Off Date", "")),
            }

            save_final_eval_record(updated)
            st.success("Teacher section submitted. Your appraiser can now complete their section.")
            _rerun()

    refreshed = get_teacher_final_eval_record(teacher_email)
    both_signed = evaluator_signed_off(teacher_email) and teacher_signed_off_final_eval(teacher_email)

    st.divider()
    st.markdown("### Appraiser Review")

    if not appraiser_final_eval_completed(teacher_email):
        st.info("Your appraiser has not completed this section yet.")

    elif not evaluator_signed_off(teacher_email):
        st.info("Your appraiser has completed the evaluation. The meeting will take place before sign-off.")

    elif evaluator_signed_off(teacher_email) and not teacher_signed_off_final_eval(teacher_email):
        st.info("The evaluation has been discussed and signed off by the appraiser. Please complete teacher sign-off after the meeting.")
        st.caption("The teacher’s signature indicates that he or she has seen and discussed the evaluation; it does not necessarily denote agreement with the report.")

        if st.button("✍️ Teacher Sign Off"):
            now_str = now_ist_str()
            refreshed["Last Edited On"] = now_str
            refreshed["Teacher Sign Off"] = "Yes"
            refreshed["Teacher Sign Off Date"] = now_str
            save_final_eval_record(refreshed)
            st.success("Teacher sign-off completed.")
            _rerun()

    if both_signed:
        st.markdown("#### Ratings on Individual Rubrics")
    
        rating_colour_map = {
            "Highly Effective": "#d4edda",
            "Effective": "#d1ecf1",
            "Improvement Necessary": "#fff3cd",
            "Does Not Meet Standards": "#f8d7da",
        }
    
        text_colour_map = {
            "Highly Effective": "#155724",
            "Effective": "#0c5460",
            "Improvement Necessary": "#856404",
            "Does Not Meet Standards": "#721c24",
        }
    
        cols = st.columns(2)
        domain_rows = final_eval_domain_rows()
    
        for i, (col_name, label) in enumerate(domain_rows):
            rating_value = safe_text(refreshed.get(col_name, ""))
            bg = rating_colour_map.get(rating_value, "#f4f4f4")
            fg = text_colour_map.get(rating_value, "#222")
    
            with cols[i % 2]:
                st.markdown(
                    f"""
                    <div style="
                        border: 1px solid #e6e6e6;
                        border-radius: 12px;
                        padding: 14px 16px;
                        margin-bottom: 12px;
                        background: #ffffff;
                        box-shadow: 0 1px 4px rgba(0,0,0,0.06);
                    ">
                        <div style="
                            font-size: 14px;
                            font-weight: 600;
                            color: #333;
                            margin-bottom: 10px;
                        ">
                            {label}
                        </div>
                        <div style="
                            display: inline-block;
                            padding: 8px 12px;
                            border-radius: 999px;
                            background: {bg};
                            color: {fg};
                            font-weight: 700;
                            font-size: 13px;
                        ">
                            {rating_value}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
    
        st.markdown("### Overall Rating")
    
        overall_rating = safe_text(refreshed.get("Overall Rating", ""))
        overall_bg = rating_colour_map.get(overall_rating, "#f4f4f4")
        overall_fg = text_colour_map.get(overall_rating, "#222")
    
        st.markdown(
            f"""
            <div style="
                border: 2px solid #dcdcdc;
                border-radius: 14px;
                padding: 18px;
                margin-top: 8px;
                margin-bottom: 14px;
                background: #fafafa;
                box-shadow: 0 1px 6px rgba(0,0,0,0.05);
            ">
                <div style="
                    font-size: 15px;
                    font-weight: 600;
                    color: #333;
                    margin-bottom: 12px;
                ">
                    Final Overall Rating
                </div>
                <div style="
                    display: inline-block;
                    padding: 10px 16px;
                    border-radius: 999px;
                    background: {overall_bg};
                    color: {overall_fg};
                    font-weight: 700;
                    font-size: 15px;
                ">
                    {overall_rating}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    
        st.markdown("### Appraiser Comments")
        st.markdown(
            f"""
            <div style="
                border: 1px solid #e6e6e6;
                border-radius: 12px;
                padding: 16px;
                background: #ffffff;
                box-shadow: 0 1px 4px rgba(0,0,0,0.06);
                line-height: 1.6;
                color: #333;
                margin-bottom: 12px;
            ">
                {safe_text(refreshed.get("Overall Comments", "")).replace("\n", "<br>")}
            </div>
            """,
            unsafe_allow_html=True
        )

        if evaluator_signed_off(teacher_email):
            st.success(f"{appraiser} signed off on {safe_text(refreshed.get('Evaluator Sign Off Date', ''))}")

        if teacher_signed_off_final_eval(teacher_email):
            st.success(f"{teacher_name} signed off on {safe_text(refreshed.get('Teacher Sign Off Date', ''))}")

        st.caption("The teacher’s signature indicates that he or she has seen and discussed the evaluation; it does not necessarily denote agreement with the report.")

# =========================
# Page: Admin Panel (Admin & Super Admin)
# =========================
if tab == "Admin" and i_am_admin:
    st.header("👩‍💼 Admin Panel")

    me = users_df[users_df["Email"] == st.session_state.auth_email].iloc[0]
    my_name = me.get("Name", st.session_state.auth_email)
    my_role = me.get("Role", "").strip().lower()
    my_first = my_name.split()[0].strip().lower()

    # Campus-awareness
    has_campus_col = "Campus" in users_df.columns
    my_campus = str(me.get("Campus", "")).strip() if has_campus_col else ""
    campus_series = (
        users_df["Campus"].astype(str).str.strip()
        if has_campus_col and my_campus
        else None
    )

    # Admins only see their assigned teachers; Super Admin sees all teachers in *their campus*
    if my_role == "sadmin":
        if campus_series is not None:
            mask = (users_df["Role"] == "user") & (campus_series == my_campus)
            assigned = users_df[mask]
            st.info(f"Super Admin access: viewing **all teachers** in the **{my_campus}** campus.")
        else:
            assigned = users_df[users_df["Role"] == "user"]
            st.info("Super Admin access: viewing **all teachers** in the school.")
    else:
        # allow multiple appraisers per teacher (comma-separated)
        def matches_appraiser(cell):
            if pd.isna(cell):
                return False
            appraisers = [a.strip().lower() for a in str(cell).split(",")]
            return my_first in appraisers

        if not users_df.empty:
            base_mask = users_df["Appraiser"].apply(matches_appraiser)
            if campus_series is not None:
                base_mask = base_mask & (campus_series == my_campus)
            assigned = users_df[base_mask]
        else:
            assigned = pd.DataFrame()

    resp_df = load_responses_df()
    if assigned.empty:
        st.info("No teachers found for your role in the Users sheet.")
    else:
        if admin_view_mode == "Summary of Teachers":
            summary_rows = []
            initial_submitted_count = 0
            final_submitted_count = 0
            total_count = len(assigned)
    
            for _, teacher in assigned.iterrows():
                teacher_email = teacher["Email"].strip().lower()
                teacher_name = teacher["Name"]
    
                submissions = resp_df[resp_df["Email"] == teacher_email] if not resp_df.empty else pd.DataFrame()
    
                # Backward compatibility: if old rows existed before Assessment Cycle was added
                if not submissions.empty:
                    if "Assessment Cycle" not in submissions.columns:
                        submissions = submissions.copy()
                        submissions["Assessment Cycle"] = "Initial"
                    else:
                        submissions = submissions.copy()
                        submissions["Assessment Cycle"] = submissions["Assessment Cycle"].replace("", "Initial")
    
                initial_subs = submissions[submissions["Assessment Cycle"] == "Initial"] if not submissions.empty else pd.DataFrame()
                final_subs = submissions[submissions["Assessment Cycle"] == "Final"] if not submissions.empty else pd.DataFrame()
    
                initial_status = "✅ Submitted" if not initial_subs.empty else "❌ Not Submitted"
                final_status = "✅ Submitted" if not final_subs.empty else "❌ Not Submitted"
    
                last_initial_date = initial_subs["Timestamp"].max() if not initial_subs.empty else "-"
                last_final_date = final_subs["Timestamp"].max() if not final_subs.empty else "-"
    
                if not initial_subs.empty:
                    initial_submitted_count += 1
                if not final_subs.empty:
                    final_submitted_count += 1
    
                summary_rows.append({
                    "Teacher": teacher_name,
                    "Email": teacher_email,
                    "Initial Status": initial_status,
                    "Final Status": final_status,
                    "Teacher Final Eval": "✅ Submitted" if teacher_final_eval_completed(teacher_email) else "❌ Pending",
                    "Appraiser Final Eval": "✅ Completed" if appraiser_final_eval_completed(teacher_email) else "❌ Pending",
                    "Last Initial": last_initial_date,
                    "Last Final": last_final_date,
                })
    
            summary_df = pd.DataFrame(summary_rows)
    
            render_metrics([
                (total_count, "Total teachers", ""),
                (initial_submitted_count, "Initial submitted", "green"),
                (final_submitted_count, "Final submitted", "green"),
                (total_count - final_submitted_count, "Final pending", "amber"),
            ])
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                pct_i = round((initial_submitted_count/total_count)*100,1) if total_count else 0
                st.markdown(f"**Initial** — {initial_submitted_count}/{total_count} ({pct_i}%)")
                st.progress(initial_submitted_count / total_count if total_count else 0)
            with col_p2:
                pct_f = round((final_submitted_count/total_count)*100,1) if total_count else 0
                st.markdown(f"**Final** — {final_submitted_count}/{total_count} ({pct_f}%)")
                st.progress(final_submitted_count / total_count if total_count else 0)
            st.markdown('<div class="sec-heading" style="margin-top:14px">Teacher status</div>', unsafe_allow_html=True)
            st.dataframe(summary_df, use_container_width=True)

        
        if admin_view_mode == "Self-Assessment Grid":
            # 🔹 Submissions Grid (My Appraisees) with color coding
            st.divider()
            st.subheader("📊 Submissions Grid (My Appraisees)")
            
            if not resp_df.empty:
                appraisee_emails = assigned["Email"].str.strip().str.lower().tolist()
                df = resp_df[resp_df["Email"].str.strip().str.lower().isin(appraisee_emails)]
            
                if not df.empty:
                    # Replace full text with acronyms
                    mapping = {
                        "Highly Effective": "HE",
                        "Effective": "E",
                        "Improvement Necessary": "IN",
                        "Does Not Meet Standards": "DNMS"
                    }
                    df = df.replace(mapping)
            
                           
                    styled_df = df.style.map(highlight_ratings, subset=df.columns[4:])
                    st.dataframe(styled_df, use_container_width=True)
                    
                    st.download_button(
                        "📥 Download My Appraisees’ Grid (CSV)",
                        data=df.to_csv(index=False).encode("utf-8"),
                        file_name=f"{st.session_state.auth_name}_appraisees_grid.csv",
                        mime="text/csv",
                    )
                else:
                    st.info("ℹ️ No rubric submissions yet from your appraisees.")

        if admin_view_mode == "View Teacher Self-Assessment":
            teacher_choice = st.selectbox("Select a teacher", assigned["Name"].tolist())

            if teacher_choice:
                teacher_email = assigned.loc[assigned["Name"] == teacher_choice, "Email"].iloc[0]
                rows = resp_df[resp_df["Email"] == teacher_email] if not resp_df.empty else pd.DataFrame()
                latest_initial, latest_final, comparison_df = build_initial_final_comparison(rows)

                has_init = latest_initial is not None and not latest_initial.empty
                has_fin = latest_final is not None and not latest_final.empty
                fe_sub = teacher_final_eval_completed(teacher_email)
                app_done = appraiser_final_eval_completed(teacher_email)
                app_signed = evaluator_signed_off(teacher_email)

                # Status summary row
                s_cols = st.columns(4)
                with s_cols[0]:
                    st.markdown(f"**Initial**<br>{'✅ Submitted' if has_init else '❌ Missing'}", unsafe_allow_html=True)
                    if has_init: st.caption(safe_text(latest_initial.iloc[0].get("Timestamp", "")))
                with s_cols[1]:
                    st.markdown(f"**Final**<br>{'✅ Submitted' if has_fin else '❌ Missing'}", unsafe_allow_html=True)
                    if has_fin: st.caption(safe_text(latest_final.iloc[0].get("Timestamp", "")))
                with s_cols[2]:
                    st.markdown(f"**Teacher final eval**<br>{'✅ Done' if fe_sub else '❌ Pending'}", unsafe_allow_html=True)
                with s_cols[3]:
                    st.markdown(f"**Appraiser section**<br>{'✅ Done' if app_done else '❌ Pending'}", unsafe_allow_html=True)

                st.divider()
                st.markdown('<div class="sec-heading">Initial vs Final Comparison</div>', unsafe_allow_html=True)

                if not has_init and not has_fin:
                    st.warning("No submissions found for this teacher.")
                elif has_init and not has_fin:
                    st.info("Only Initial submission exists. Comparison will appear once Final is submitted.")
                    init_display = latest_initial.copy().replace({
                        "Highly Effective": "HE", "Effective": "E",
                        "Improvement Necessary": "IN", "Does Not Meet Standards": "DNMS"
                    })
                    st.dataframe(init_display.style.map(highlight_ratings, subset=init_display.columns[5:]), use_container_width=True)
                else:
                    if not comparison_df.empty:
                        display_df = comparison_df[["Domain", "Strand", "Explanation", "Initial", "Final", "Trend"]].copy()
                        render_comparison_table_new(display_df)
                        appraiser_name = safe_text(rows.sort_values("Timestamp", ascending=False).head(1).iloc[0].get("Appraiser", ""))
                        printable_html = build_printable_comparison_html(
                            teacher_name=teacher_choice, teacher_email=teacher_email,
                            appraiser=appraiser_name, latest_initial=latest_initial,
                            latest_final=latest_final, display_df=display_df
                        )

                st.divider()
                        
                if rows.empty:
                    st.warning(f"No submission found for {teacher_choice}.")
                else:    
                    
                    st.subheader("Final Evaluation")
                    fe_record = get_teacher_final_eval_record(teacher_email)
        
                    if not teacher_final_eval_completed(teacher_email):
                        st.info("Teacher has not yet submitted their Final Evaluation section.")
                    else:
                        st.success("Teacher section submitted.")
        
                        st.write(f"**Subject Area:** {safe_text(fe_record.get('Subject Area', ''))}")
                        st.write("**Student Survey Feedback:**")
                        st.write(safe_text(fe_record.get("Student Survey Feedback", "")))
                        st.write("**Overall Reflection:**")
                        st.write(safe_text(fe_record.get("Overall Reflection", "")))
        
                        if teacher_signed_off_final_eval(teacher_email):
                            st.divider()
                            render_final_evaluation_review_panel(fe_record, heading="Final Signed-Off Review")
        
                            if evaluator_signed_off(teacher_email):
                                st.success(f"{st.session_state.auth_name} signed off on {safe_text(fe_record.get('Evaluator Sign Off Date', ''))}")
        
                            if teacher_signed_off_final_eval(teacher_email):
                                st.success(f"{teacher_choice} signed off on {safe_text(fe_record.get('Teacher Sign Off Date', ''))}")
                            
                                final_doc_record = fe_record.copy()
                                final_doc_record["Teacher Name"] = teacher_choice
                                final_doc_record["Appraiser"] = title_case_name(final_doc_record.get("Appraiser", st.session_state.auth_name))
    
                                final_docx = generate_final_evaluation_docx(final_doc_record)
    
                                st.download_button(
                                    "📄 Download Final Evaluation Summary (DOCX)",
                                    data=final_docx,
                                    file_name=f"{teacher_choice}_final_evaluation_summary.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"{teacher_email}_final_eval_docx"
                                )
                            
                        else:
                            appraiser_locked = (
                                not is_before_deadline(FINAL_EVAL_APPRAISER_DEADLINE)
                                or teacher_signed_off_final_eval(teacher_email)
                            )
        
                            st.caption(f"Appraiser deadline (IST): {FINAL_EVAL_APPRAISER_DEADLINE.strftime('%d %b %Y, %I:%M %p')}")
        
                            domain_values = {}
                            for rating_col, label in final_eval_domain_rows():
                                existing = safe_text(fe_record.get(rating_col, ""))
        
                                default_index = FINAL_EVAL_RATINGS.index(existing) if existing in FINAL_EVAL_RATINGS else 0
        
                                domain_values[rating_col] = st.selectbox(
                                    label,
                                    FINAL_EVAL_RATINGS,
                                    index=default_index,
                                    disabled=appraiser_locked,
                                    key=f"{teacher_email}_{rating_col}"
                                )
        
                            existing_overall = safe_text(fe_record.get("Overall Rating", ""))
                            default_overall_index = FINAL_EVAL_RATINGS.index(existing_overall) if existing_overall in FINAL_EVAL_RATINGS else 0
            
                            overall_rating = st.selectbox(
                                    "Overall Rating",
                                    FINAL_EVAL_RATINGS,
                                    index=default_overall_index,
                                    disabled=appraiser_locked,
                                    key=f"{teacher_email}_overall_rating"
                                )
        
                            overall_comments = st.text_area(
                                "Overall Comments (150 words or less)",
                                value=safe_text(fe_record.get("Overall Comments", "")),
                                height=180,
                                disabled=appraiser_locked,
                                key=f"{teacher_email}_overall_comments"
                            )
        
                            comments_wc = count_words(overall_comments)
                            st.caption(f"Word count: {comments_wc}/{FINAL_EVAL_MAX_WORDS_COMMENTS}")
        
                            col_a, col_b = st.columns(2)
        
                            with col_a:
                                if st.button(
                                    "💾 Save Appraiser Section",
                                    disabled=appraiser_locked or comments_wc > FINAL_EVAL_MAX_WORDS_COMMENTS,
                                    key=f"{teacher_email}_save_appraiser_eval"
                                ):
                                    now_str = now_ist_str()
                                    updated = fe_record.copy()
                                    updated["Last Edited On"] = now_str
                                    updated["Appraiser Started"] = "Yes"
        
                                    for k, v in domain_values.items():
                                        updated[k] = v
        
                                    updated["Overall Rating"] = overall_rating
                                    updated["Overall Comments"] = overall_comments
        
                                    save_final_eval_record(updated)
                                    st.success("Appraiser section saved.")
                                    _rerun()
        
                            with col_b:
                                if st.button(
                                    "✅ Submit Appraiser Section",
                                    disabled=appraiser_locked or comments_wc > FINAL_EVAL_MAX_WORDS_COMMENTS,
                                    key=f"{teacher_email}_submit_appraiser_eval"
                                ):
                                    now_str = now_ist_str()
                                    updated = fe_record.copy()
                                    updated["Last Edited On"] = now_str
                                    updated["Appraiser Started"] = "Yes"
                                    updated["Appraiser Completed"] = "Yes"
                                    updated["Appraiser Completed On"] = now_str
        
                                    for k, v in domain_values.items():
                                        updated[k] = v
        
                                    updated["Overall Rating"] = overall_rating
                                    updated["Overall Comments"] = overall_comments
        
                                    save_final_eval_record(updated)
                                    st.success("Appraiser section submitted.")
                                    _rerun()
        
                            refreshed_fe = get_teacher_final_eval_record(teacher_email)
        
                        if appraiser_final_eval_completed(teacher_email) and not evaluator_signed_off(teacher_email) and not appraiser_locked:
                            st.info("Only click Appraiser Sign Off after the evaluation has been discussed with the teacher in the meeting. Once signed off, the evaluation will become visible to the teacher for final teacher sign-off.")
    
                            if st.button("✍️ Appraiser Sign Off", key=f"{teacher_email}_evaluator_signoff"):
                                now_str = now_ist_str()
                                refreshed_fe["Last Edited On"] = now_str
                                refreshed_fe["Evaluator Sign Off"] = "Yes"
                                refreshed_fe["Evaluator Sign Off Date"] = now_str
                                save_final_eval_record(refreshed_fe)
                                st.success("Appraiser sign-off completed.")
                                _rerun()
        
                            if evaluator_signed_off(teacher_email):
                                st.success(f"{st.session_state.auth_name} signed off on {safe_text(refreshed_fe.get('Evaluator Sign Off Date', ''))}")


# =========================
# Page: Super Admin Panel
# =========================
if tab == "Super Admin" and i_am_sadmin:
    st.header("🏫 Super Admin Panel — Campus View")

    my_campus = str(st.session_state.get("auth_campus", "") or "").strip()
    has_campus_col = "Campus" in users_df.columns
    campus_series = (
        users_df["Campus"].astype(str).str.strip()
        if has_campus_col and my_campus
        else None
    )

    if campus_series is not None:
        assigned = users_df[(users_df["Role"] == "user") & (campus_series == my_campus)]
    else:
        assigned = users_df[users_df["Role"] == "user"]

    resp_df = load_responses_df()

    if assigned.empty:
        st.info("No teachers found for this campus.")
    else:
        summary_rows = []
        initial_submitted_count = 0
        final_submitted_count = 0
        total_count = len(assigned)

        for _, teacher in assigned.iterrows():
            teacher_email = teacher["Email"].strip().lower()
            teacher_name = teacher["Name"]

            submissions = resp_df[resp_df["Email"] == teacher_email] if not resp_df.empty else pd.DataFrame()

            if not submissions.empty:
                if "Assessment Cycle" not in submissions.columns:
                    submissions = submissions.copy()
                    submissions["Assessment Cycle"] = "Initial"
                else:
                    submissions = submissions.copy()
                    submissions["Assessment Cycle"] = submissions["Assessment Cycle"].replace("", "Initial")

            initial_subs = submissions[submissions["Assessment Cycle"] == "Initial"] if not submissions.empty else pd.DataFrame()
            final_subs = submissions[submissions["Assessment Cycle"] == "Final"] if not submissions.empty else pd.DataFrame()

            initial_status = "✅ Submitted" if not initial_subs.empty else "❌ Not Submitted"
            final_status = "✅ Submitted" if not final_subs.empty else "❌ Not Submitted"

            last_initial_date = initial_subs["Timestamp"].max() if not initial_subs.empty else "-"
            last_final_date = final_subs["Timestamp"].max() if not final_subs.empty else "-"

            if not initial_subs.empty:
                initial_submitted_count += 1
            if not final_subs.empty:
                final_submitted_count += 1

            summary_rows.append({
                "Teacher": teacher_name,
                "Email": teacher_email,
                "Initial Status": initial_status,
                "Final Status": final_status,
                "Teacher Final Eval": "✅ Submitted" if teacher_final_eval_completed(teacher_email) else "❌ Pending",
                "Appraiser Final Eval": "✅ Completed" if appraiser_final_eval_completed(teacher_email) else "❌ Pending",
                "Last Initial": last_initial_date,
                "Last Final": last_final_date,
            })

        summary_df = pd.DataFrame(summary_rows)

        if sadmin_view_mode == "Summary of Teachers":
            render_metrics([
                (total_count, "Total teachers", ""),
                (initial_submitted_count, "Initial submitted", "green"),
                (final_submitted_count, "Final submitted", "green"),
                (total_count - final_submitted_count, "Final pending", "amber"),
            ])
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                pct_i = round((initial_submitted_count/total_count)*100,1) if total_count else 0
                st.markdown(f"**Initial** — {initial_submitted_count}/{total_count} ({pct_i}%)")
                st.progress(initial_submitted_count / total_count if total_count else 0)
            with col_p2:
                pct_f = round((final_submitted_count/total_count)*100,1) if total_count else 0
                st.markdown(f"**Final** — {final_submitted_count}/{total_count} ({pct_f}%)")
                st.progress(final_submitted_count / total_count if total_count else 0)
            st.markdown('<div class="sec-heading" style="margin-top:14px">All teachers</div>', unsafe_allow_html=True)
            st.dataframe(summary_df, use_container_width=True)

        if sadmin_view_mode == "Self-Assessment Grid":
            st.divider()
            st.subheader("📊 Submissions Grid (Campus)")

            if not resp_df.empty:
                teacher_emails = assigned["Email"].str.strip().str.lower().tolist()
                df = resp_df[resp_df["Email"].str.strip().str.lower().isin(teacher_emails)]

                if not df.empty:
                    mapping = {
                        "Highly Effective": "HE",
                        "Effective": "E",
                        "Improvement Necessary": "IN",
                        "Does Not Meet Standards": "DNMS"
                    }
                    df = df.replace(mapping)

                    styled_df = df.style.map(highlight_ratings, subset=df.columns[4:])
                    st.dataframe(styled_df, use_container_width=True)

                    st.download_button(
                        "📥 Download Campus Grid (CSV)",
                        data=df.to_csv(index=False).encode("utf-8"),
                        file_name=f"{my_campus or 'campus'}_submissions_grid.csv",
                        mime="text/csv",
                    )
                else:
                    st.info("ℹ️ No rubric submissions yet for this campus.")

        if sadmin_view_mode == "View Teacher Self-Assessment":
            st.divider()
            st.subheader("🔎 View Individual Submissions")

            teacher_choice = st.selectbox(
                "Select a teacher",
                assigned["Name"].tolist(),
                key="sadmin_teacher_choice"
            )

            if teacher_choice:
                teacher_email = assigned.loc[assigned["Name"] == teacher_choice, "Email"].iloc[0]
                rows = resp_df[resp_df["Email"] == teacher_email] if not resp_df.empty else pd.DataFrame()

                latest_initial, latest_final, comparison_df = build_initial_final_comparison(rows)

                st.subheader(f"Initial vs Final Comparison for {teacher_choice}")

                col1, col2 = st.columns(2)

                with col1:
                    if latest_initial is not None and not latest_initial.empty:
                        st.info(f"Initial submitted: {safe_text(latest_initial.iloc[0].get('Timestamp', ''))}")
                    else:
                        st.warning("No Initial submission found.")

                with col2:
                    if latest_final is not None and not latest_final.empty:
                        st.info(f"Final submitted: {safe_text(latest_final.iloc[0].get('Timestamp', ''))}")
                    else:
                        st.warning("No Final submission found.")

                if not comparison_df.empty:
                    display_df = comparison_df[["Domain", "Strand", "Explanation", "Initial", "Final", "Trend"]].copy()
                    render_comparison_table_new(display_df)

                st.divider()

                if rows.empty:
                    st.warning(f"No submission found for {teacher_choice}.")
                else:
                    st.subheader("Final Evaluation")
                    fe_record = get_teacher_final_eval_record(teacher_email)

                    if not teacher_final_eval_completed(teacher_email):
                        st.info("Teacher has not yet submitted their Final Evaluation section.")
                    else:
                        st.success("Teacher section submitted.")

                        st.write(f"**Subject Area:** {safe_text(fe_record.get('Subject Area', ''))}")
                        st.write("**Student Survey Feedback:**")
                        st.write(safe_text(fe_record.get("Student Survey Feedback", "")))
                        st.write("**Overall Reflection:**")
                        st.write(safe_text(fe_record.get("Overall Reflection", "")))

                        if teacher_signed_off_final_eval(teacher_email):
                            st.divider()
                            render_final_evaluation_review_panel(fe_record, heading="Final Signed-Off Review")

                            if evaluator_signed_off(teacher_email):
                                st.success(f"{safe_text(fe_record.get('Appraiser', st.session_state.auth_name))} signed off on {safe_text(fe_record.get('Evaluator Sign Off Date', ''))}")

                            if teacher_signed_off_final_eval(teacher_email):
                                st.success(f"{teacher_choice} signed off on {safe_text(fe_record.get('Teacher Sign Off Date', ''))}")

                            final_doc_record = fe_record.copy()
                            final_doc_record["Teacher Name"] = teacher_choice
                            final_doc_record["Appraiser"] = title_case_name(final_doc_record.get("Appraiser", st.session_state.auth_name))

                            final_docx = generate_final_evaluation_docx(final_doc_record)

                            st.download_button(
                                "📄 Download Final Evaluation Summary (DOCX)",
                                data=final_docx,
                                file_name=f"{teacher_choice}_final_evaluation_summary.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"{teacher_email}_sadmin_final_eval_docx"
                            )
                        else:
                            appraiser_locked = (
                                not is_before_deadline(FINAL_EVAL_APPRAISER_DEADLINE)
                                or teacher_signed_off_final_eval(teacher_email)
                            )

                            st.caption(f"Appraiser deadline (IST): {FINAL_EVAL_APPRAISER_DEADLINE.strftime('%d %b %Y, %I:%M %p')}")

                            domain_values = {}
                            for rating_col, label in final_eval_domain_rows():
                                existing = safe_text(fe_record.get(rating_col, ""))
                                default_index = FINAL_EVAL_RATINGS.index(existing) if existing in FINAL_EVAL_RATINGS else 0

                                domain_values[rating_col] = st.selectbox(
                                    label,
                                    FINAL_EVAL_RATINGS,
                                    index=default_index,
                                    disabled=appraiser_locked,
                                    key=f"{teacher_email}_sadmin_{rating_col}"
                                )

                            existing_overall = safe_text(fe_record.get("Overall Rating", ""))
                            default_overall_index = FINAL_EVAL_RATINGS.index(existing_overall) if existing_overall in FINAL_EVAL_RATINGS else 0

                            overall_rating = st.selectbox(
                                "Overall Rating",
                                FINAL_EVAL_RATINGS,
                                index=default_overall_index,
                                disabled=appraiser_locked,
                                key=f"{teacher_email}_sadmin_overall_rating"
                            )

                            overall_comments = st.text_area(
                                "Overall Comments (150 words or less)",
                                value=safe_text(fe_record.get("Overall Comments", "")),
                                height=180,
                                disabled=appraiser_locked,
                                key=f"{teacher_email}_sadmin_overall_comments"
                            )

                            comments_wc = count_words(overall_comments)
                            st.caption(f"Word count: {comments_wc}/{FINAL_EVAL_MAX_WORDS_COMMENTS}")

                            col_a, col_b = st.columns(2)

                            with col_a:
                                if st.button(
                                    "💾 Save Appraiser Section",
                                    disabled=appraiser_locked or comments_wc > FINAL_EVAL_MAX_WORDS_COMMENTS,
                                    key=f"{teacher_email}_sadmin_save_appraiser_eval"
                                ):
                                    now_str = now_ist_str()
                                    updated = fe_record.copy()
                                    updated["Last Edited On"] = now_str
                                    updated["Appraiser Started"] = "Yes"

                                    for k, v in domain_values.items():
                                        updated[k] = v

                                    updated["Overall Rating"] = overall_rating
                                    updated["Overall Comments"] = overall_comments

                                    save_final_eval_record(updated)
                                    st.success("Appraiser section saved.")
                                    _rerun()

                            with col_b:
                                if st.button(
                                    "✅ Submit Appraiser Section",
                                    disabled=appraiser_locked or comments_wc > FINAL_EVAL_MAX_WORDS_COMMENTS,
                                    key=f"{teacher_email}_sadmin_submit_appraiser_eval"
                                ):
                                    now_str = now_ist_str()
                                    updated = fe_record.copy()
                                    updated["Last Edited On"] = now_str
                                    updated["Appraiser Started"] = "Yes"
                                    updated["Appraiser Completed"] = "Yes"
                                    updated["Appraiser Completed On"] = now_str

                                    for k, v in domain_values.items():
                                        updated[k] = v

                                    updated["Overall Rating"] = overall_rating
                                    updated["Overall Comments"] = overall_comments

                                    save_final_eval_record(updated)
                                    st.success("Appraiser section submitted.")
                                    _rerun()

                            refreshed_fe = get_teacher_final_eval_record(teacher_email)

                            if appraiser_final_eval_completed(teacher_email) and not evaluator_signed_off(teacher_email) and not appraiser_locked:
                                st.info("Only click Appraiser Sign Off after the evaluation has been discussed with the teacher in the meeting. Once signed off, the evaluation will become visible to the teacher for final teacher sign-off.")

                                if st.button("✍️ Appraiser Sign Off", key=f"{teacher_email}_sadmin_evaluator_signoff"):
                                    now_str = now_ist_str()
                                    refreshed_fe["Last Edited On"] = now_str
                                    refreshed_fe["Evaluator Sign Off"] = "Yes"
                                    refreshed_fe["Evaluator Sign Off Date"] = now_str
                                    save_final_eval_record(refreshed_fe)
                                    st.success("Appraiser sign-off completed.")
                                    _rerun()

                            if evaluator_signed_off(teacher_email):
                                st.success(f"{safe_text(fe_record.get('Appraiser', st.session_state.auth_name))} signed off on {safe_text(refreshed_fe.get('Evaluator Sign Off Date', ''))}")
